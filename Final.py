import os
import re
import pdfplumber
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# ===== FUNCIONES =====

def extract_data_from_pdf(pdf_path):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())

        def extract_field(pattern):
            match = re.search(pattern, text)
            return match.group(1).strip() if match else ""

        data = {
            '[Nombre]': extract_field(r'Nombre o Razón Social:\s*([^\n]+)'),
            '[Rut]': extract_field(r'Rut:\s*([^\s]+)'),
            '[Giro]': extract_field(r'Giro:\s*([^\n]+)'),
            '[Direccion]': extract_field(r'Dirección Comercial:\s*([^\n]+)'),
            '[Ciudad]': extract_field(r'Ciudad:\s*([^\n]+)'),
            '[Contacto]': extract_field(r'Contacto:\s*([^\n]+?)(?=\s*Proyecto Asociado:|\n|$)'),
            '[Email]': extract_field(r'e-mail:\s*([^\s]+)'),
            '[Proyecto]': extract_field(r'Proyecto Asociado:\s*([^\n]+)')
        }

        return data

    except Exception as e:
        print(f"Error al procesar PDF: {e}")
        return None


def fill_lab_report(doc, data):
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace(key, value)


def insert_element_after(paragraph, element):
    """Inserta un elemento (tabla, párrafo) justo después de un párrafo"""
    p = paragraph._p
    p.addnext(element._element)


def insert_excel_table(doc, excel_path):
    df = pd.read_excel(excel_path)

    # Buscar el párrafo con "Resultados"
    for i, paragraph in enumerate(doc.paragraphs):
        if "Resultados" in paragraph.text:
            insert_index = i
            break
    else:
        insert_index = len(doc.paragraphs)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr_cells[j].text = str(col)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for j, value in enumerate(row):
            row_cells[j].text = str(value)

    insert_element_after(doc.paragraphs[insert_index], table)


# ===== PROCESO PRINCIPAL =====

def generar_informe(pdf_path, plantilla_path, excel_path, output_path):
    data = extract_data_from_pdf(pdf_path)
    if not data:
        print("Error: no se pudo extraer información del PDF.")
        return

    doc = Document(plantilla_path)

    # Rellenar con datos del PDF
    fill_lab_report(doc, data)

    # Insertar tabla desde Excel
    insert_excel_table(doc, excel_path)

    doc.save(output_path)
    print(f"Documento generado: {output_path}")


# ===== LLAMADA DE EJEMPLO =====

if __name__ == "__main__":
    pdf = "uploads/cliente.pdf"
    plantilla = "template.docx"
    excel = "uploads/data.xlsx"
    salida = "documento_final.docx"

    generar_informe(pdf, plantilla, excel, salida)
