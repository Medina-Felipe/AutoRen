import pdfplumber
import re
from docx import Document
import os
import pandas as pd
from docxtpl import DocxTemplate

def extract_data_from_pdf(pdf_path):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())

        def extract_field(pattern):
            match = re.search(pattern, text)
            return match.group(1).strip() if match else ""

        data = {
            '[Nombre]': extract_field(r'Nombre o Razón Social:\s*([^\n]+)'),
            '[Rut]': extract_field(r'Rut:\s*([^\s]+)'),
            '[Giro]': extract_field(r'Giro:\s*([^\n]+)'),
            '[Direccion]': extract_field(r'Dirección Comercial:\s*([^\n]+)'),
            '[Ciudad]': extract_field(r'Ciudad:\s*([^\n]+)'),
            '[Contacto]': extract_field(r'Contacto:\s*([^\n]+?)(?=\s*Proyecto Asociado:|\n|$)'),
            '[Email]': extract_field(r'e-mail:\s*([^\s]+)'),
            '[Proyecto]': extract_field(r'Proyecto Asociado:\s*([^\n]+)')
            
        }

        return data

    except Exception as e:
        print(f"Error al procesar PDF: {e}")
        return None

def fill_lab_report(template_path, output_path, data):
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace(key, value)

    doc.save(output_path)

def process_single_order(pdf_path, template_path):
    data = extract_data_from_pdf(pdf_path)
    if not data:
        return False, None

    nombre_seguro = data.get('[Nombre]', 'sin_nombre').replace(' ', '_')
    output_path = f"output/informe_{nombre_seguro}.docx"

    try:
        fill_lab_report(template_path, output_path, data)
        return True, output_path
    except Exception as e:
        print(f"Error generando informe: {e}")
        return False, None


import pandas as pd
from docx import Document


excel_path = "static/files/data_autoren.xlsx"  

df = pd.read_excel(excel_path)


doc = Document("static/files/Plantilla.docx")


doc.add_page_break()
doc.add_paragraph("Datos Adicionales desde Excel", style='Heading 1')


table = doc.add_table(rows=1, cols=len(df.columns))
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
for i, col_name in enumerate(df.columns):
    hdr_cells[i].text = str(col_name)


for _, row in df.iterrows():
    row_cells = table.add_row().cells
    for i, item in enumerate(row):
        row_cells[i].text = str(item)

doc.save("documento_final.docx")