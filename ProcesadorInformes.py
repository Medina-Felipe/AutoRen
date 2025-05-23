import pdfplumber
import re
from docx import Document
import os

def extract_data_from_pdf(pdf_path):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            # Leer todas las páginas
            text = ""
            for page in pdf.pages:
                text += page.extract_text() + "\n"
            
            # Función para extraer datos usando expresiones regulares
            def extract_field(pattern):
                match = re.search(pattern, text)
                return match.group(1).strip() if match else None

            # Extraer los datos
            data = {
                'nombre': extract_field(r'Nombre o Razón Social:\s*([^\n]+)'),
                'rut': extract_field(r'Rut:\s*([^\s]+)'),
                'giro': extract_field(r'Giro:\s*([^\n]+)'),
                'direccion': extract_field(r'Dirección Comercial:\s*([^\n]+)'),
                'ciudad': extract_field(r'Ciudad:\s*([^\n]+)'),
                'contacto': extract_field(r'Contacto:\s*([^\n]+?)(?=\s*Proyecto Asociado:|\n|$)'),
                'email': extract_field(r'e-mail:\s*([^\s]+)'),
                'proyecto': extract_field(r'Proyecto Asociado:\s*([^\n]+)')
            }
            
            return data

    except Exception as e:
        print(f"Error al procesar el PDF: {str(e)}")
        return None

def fill_lab_report(template_path, output_path, data):
    doc = Document(template_path)

    # Reemplazar en párrafos generales
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

    # Reemplazar en celdas de tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace(key, value)

    doc.save(output_path)

def process_single_order(pdf_path, template_path):
    # Extraer datos del PDF
    data = extract_data_from_pdf(pdf_path)
    
    if not data:
        print("No se pudieron extraer los datos del PDF")
        return False
    
    # Crear directorio de salida si no existe
    os.makedirs("output", exist_ok=True)
    
    # Preparar los datos para la plantilla
    template_data = {
        '[Nombre]': data['nombre'],
        '[Rut]': data['rut'],
        '[Giro]': data['giro'],
        '[Direccion]': data['direccion'],
        '[Ciudad]': data['ciudad'],
        '[Contacto]': data['contacto'],
        '[Email]': data['email'],
        '[Proyecto]': data['proyecto']
    }
    
    # Generar el nombre del archivo de salida basado en el nombre del cliente
    output_filename = f"output/informe_{data['nombre'].replace(' ', '_')}.docx"
    
    # Generar el informe
    try:
        fill_lab_report(template_path, output_filename, template_data)
        print(f"Informe generado exitosamente: {output_filename}")
        return True
    except Exception as e:
        print(f"Error al generar el informe: {str(e)}")
        return False

if __name__ == "__main__":
    pdf_path = "Orden de Ingreso.pdf"
    template_path = "template.docx"
    
    print("Procesando orden de ingreso...")
    if process_single_order(pdf_path, template_path):
        print("Proceso completado exitosamente")
    else:
        print("El proceso no se completó correctamente") 