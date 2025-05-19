import pandas as pd
from docx import Document
import os

def fill_invitation(template_path, output_path, data):
    doc = Document(template_path)

    # Reemplazar en párrafos generales
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

    # Reemplazar en celdas de tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace(key, value)

    doc.save(output_path)

def generate_invitation_from_csv(csv_path, template_path):
    df = pd.read_csv(csv_path)
    os.makedirs("output", exist_ok=True)
    
    for idx, row in df.iterrows():
        data = {
            '[Nombre]': row['nombre'],
            '[Rut]': row['rut'],
            '[Giro]': row['giro'],
            '[Direccion]': row['direccion'],
            '[Ciudad]': row['ciudad'],
            '[Contacto]': row['contacto'],
            '[Email]': row['email'],
            '[Proyecto]': row['proyecto'],
        }
        output_path = f"output/invitation_{idx+1}.docx"
        fill_invitation(template_path, output_path, data)

if __name__ == "__main__":
    csv_path = 'contacts.csv'
    template_path = 'template.docx'
    generate_invitation_from_csv(csv_path, template_path)