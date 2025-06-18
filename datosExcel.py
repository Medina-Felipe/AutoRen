import pandas as pd
from docx import Document

doc = Document("template.docx")
df = pd.read_excel("uploads/data.xlsx")

for i, paragraph in enumerate(doc.paragraphs):
    if "Resultados" in paragraph.text:
        insert_index = i
        break
else:
    insert_index = len(doc.paragraphs)

table = doc.add_table(rows=1, cols=len(df.columns))
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
for j, col in enumerate(df.columns):
    hdr_cells[j].text = str(col)

for _, row in df.iterrows():
    row_cells = table.add_row().cells
    for j, value in enumerate(row):
        row_cells[j].text = str(value)

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

def insert_element_after(paragraph, element):
    """Inserta un elemento XML (tabla, párrafo) justo después del párrafo dado"""
    p = paragraph._p
    p.addnext(element._element)

insert_element_after(doc.paragraphs[insert_index], table)

doc.save("documento_final.docx")